"""
Core file routing logic. Processes a single file sequentially.
Returns a Markdown string.

File type → processing path:
  PDF                  → Gemini File API → Prompt A
  Images (jpg/png/...) → Gemini File API → Prompt B
  Audio (webm/mp3/...) → Gemini File API → Prompt C
  DOCX/DOC             → python-docx native extraction → Prompt A (text)
  XLSX/CSV             → openpyxl/pandas native → Prompt D (text)
"""

import io
from docx import Document
import pandas as pd

from services import gemini_client

# MIME type map for Gemini File API uploads
_MIME_MAP = {
    "pdf": "application/pdf",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "png": "image/png",
    "webp": "image/webp",
    "heic": "image/heic",
    "gif": "image/gif",
    "webm": "audio/webm",
    "mp3": "audio/mp3",
    "wav": "audio/wav",
    "ogg": "audio/ogg",
    "m4a": "audio/m4a",
    "aac": "audio/aac",
}

_IMAGE_EXTS = {"jpg", "jpeg", "png", "webp", "heic", "gif"}
_AUDIO_EXTS = {"webm", "mp3", "wav", "ogg", "m4a", "aac"}
_DOCX_EXTS = {"doc", "docx"}
_SHEET_EXTS = {"xlsx", "xls", "csv"}


def _ext(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


async def process_file(filename: str, data: bytes) -> str:
    """Route file to the correct Gemini prompt. Returns Markdown."""
    ext = _ext(filename)

    if ext == "pdf":
        return await _process_via_file_api(filename, data, "pdf")

    if ext in _IMAGE_EXTS:
        return await _process_image(filename, data, ext)

    if ext in _AUDIO_EXTS:
        return await _process_audio(filename, data, ext)

    if ext in _DOCX_EXTS:
        return await _process_docx(data)

    if ext in _SHEET_EXTS:
        return await _process_sheet(filename, data, ext)

    # Fallback: treat as plain text → Prompt A
    text = data.decode("utf-8", errors="replace")
    return await gemini_client.process_with_prompt_a_text(text)


# ─── Internal helpers ─────────────────────────────────────────────────────────

async def _process_via_file_api(filename: str, data: bytes, ext: str) -> str:
    mime = _MIME_MAP.get(ext, "application/octet-stream")
    uri = await gemini_client.upload_file(data, filename, mime)
    try:
        return await gemini_client.process_with_prompt_a(uri, mime)
    finally:
        await gemini_client.delete_file(uri)


async def _process_image(filename: str, data: bytes, ext: str) -> str:
    mime = _MIME_MAP.get(ext, "image/jpeg")
    uri = await gemini_client.upload_file(data, filename, mime)
    try:
        return await gemini_client.process_with_prompt_b(uri, mime)
    finally:
        await gemini_client.delete_file(uri)


async def _process_audio(filename: str, data: bytes, ext: str) -> str:
    mime = _MIME_MAP.get(ext, "audio/webm")
    uri = await gemini_client.upload_file(data, filename, mime)
    try:
        return await gemini_client.process_with_prompt_c(uri, mime)
    finally:
        await gemini_client.delete_file(uri)


async def _process_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also grab table cells
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    extracted = "\n\n".join(paragraphs)
    return await gemini_client.process_with_prompt_a_text(extracted)


async def _process_sheet(filename: str, data: bytes, ext: str) -> str:
    if ext == "csv":
        df = pd.read_csv(io.BytesIO(data))
    else:
        df = pd.read_excel(io.BytesIO(data))
    table_text = df.to_markdown(index=False) if hasattr(df, "to_markdown") else df.to_string(index=False)
    return await gemini_client.process_with_prompt_d(table_text)
